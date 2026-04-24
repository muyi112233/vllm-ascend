import sys
import os
import re
from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r'C:\Repo\vllm-ascend\docs\source\tutorials\features\GLM-5.1昇腾支持一本通-vllm-ascend v0.18.0版本配套-v1.0.docx'
doc = Document(doc_path)

output_lines = []

def get_heading_level(style_name):
    if style_name == 'Heading 1': return 1
    if style_name == 'Heading 2': return 2
    if style_name == 'Heading 3': return 3
    if style_name == 'Heading 4': return 4
    if style_name == 'Heading 5': return 5
    if style_name == 'Heading 6': return 6
    return 0

def is_code_style(style_name):
    code_styles = ['Code', 'Source Code', 'No Spacing', 'Normal (Web)']
    return style_name in code_styles

def detect_code_block(paras, start_idx):
    consecutive_code = 0
    for j in range(start_idx, min(start_idx + 3, len(paras))):
        p = paras[j]
        text = p.text.strip()
        style = p.style.name if p.style else ''
        if not text:
            continue
        if is_code_style(style):
            consecutive_code += 1
            continue
        if any(keyword in text for keyword in ['export ', 'import ', 'def ', 'class ', 'for ', 'if ', 'vllm ', 'docker ', 'pip ', 'bash ', 'python ', 'hccn_tool', 'npu-smi', 'cat ', '# ', '#!/', '--', 'nic_name=', 'local_ip=', 'node0_ip=']):
            consecutive_code += 1
            continue
        break
    return consecutive_code >= 2

in_code_block = False
code_lang = ''
prev_was_heading = False

all_paras = list(doc.paragraphs)

i = 0
while i < len(all_paras):
    para = all_paras[i]
    style_name = para.style.name if para.style else 'None'
    text = para.text.strip()

    if style_name.startswith('toc'):
        i += 1
        continue

    if style_name.startswith('Table'):
        i += 1
        continue

    if text == '目  录' or text == '目录':
        i += 1
        continue

    heading_level = get_heading_level(style_name)

    if heading_level > 0:
        if in_code_block:
            output_lines.append('```')
            output_lines.append('')
            in_code_block = False
        prefix = '#' * heading_level
        clean_text = re.sub(r'^\d+\.\d+\.\d+\.\d+\s+', '', text)
        clean_text = re.sub(r'^\d+\.\d+\.\d+\s+', '', clean_text)
        clean_text = re.sub(r'^\d+\.\d+\s+', '', clean_text)
        clean_text = re.sub(r'^\d+\s+', '', clean_text)
        output_lines.append(f'{prefix} {clean_text}')
        output_lines.append('')
        prev_was_heading = True
        i += 1
        continue

    if not text:
        if in_code_block:
            pass
        else:
            if not prev_was_heading:
                output_lines.append('')
        i += 1
        continue

    prev_was_heading = False

    is_list = style_name in ['List Paragraph', 'List Bullet', 'List Number', 'List Continue', 'List 2', 'List 3', 'List 4', 'List 5']

    if is_list:
        if in_code_block:
            output_lines.append('```')
            output_lines.append('')
            in_code_block = False
        level = 0
        pf = para.paragraph_format
        if pf and pf.left_indent is not None:
            level = int(pf.left_indent / 720)
        indent = '  ' * level
        clean_text = re.sub(r'^[\d]+\.\s*', '', text)
        output_lines.append(f'{indent}- {clean_text}')
        i += 1
        continue

    code_indicators = [
        'export ', 'import ', 'from ', 'def ', 'class ', 'for ', 'if __',
        'vllm serve', 'vllm ', 'docker run', 'docker pull', 'docker load',
        'docker exec', 'pip install', 'bash ', 'python ', 'hccn_tool',
        'npu-smi', 'cat /', '#!/', 'nic_name=', 'local_ip=', 'node0_ip=',
        '--host', '--port', '--data-parallel', '--tensor-parallel',
        '--quantization', '--seed', '--served-model', '--max-num',
        '--trust-remote', '--gpu-memory', '--enable-', '--async-',
        '--compilation-config', '--speculative-config', '--additional-config',
        '--kv-transfer-config', '--reasoning-parser', '--tool-call-parser',
        '--enforce-eager', 'IMAGES_ID=', 'NAME=', 'if [ $#',
        'if ! docker', 'echo ', 'exit ', 'fi', 'done', 'done;',
        'ASCEND_RT_VISIBLE_DEVICES', 'LD_LIBRARY_PATH',
        'PYTORCH_NPU_ALLOC_CONF', 'HCCL_', 'OMP_', 'VLLM_',
        'ASCEND_', 'GLOO_SOCKET', 'TP_SOCKET', 'TASK_QUEUE',
        'ACL_OP_INIT', 'VLLM_NIXL', 'ASCEND_RT',
        'self.', 'return ', 'parser.add_argument', 'args =',
        'command =', 'subprocess.', 'process.', 'async ',
        'await ', 'httpx.', 'fastapi', 'heapq.', 'logger.',
        'asyncio.', 'threading.', 'multiprocessing.',
        'response =', 'headers =', 'req_data', 'client.',
        'app =', 'proxy_state', 'global_args',
        'sys.exit', 'os.path', 'os.environ',
        'kv_connector', 'kv_role', 'kv_port', 'engine_id',
        'kv_connector_extra_config', 'use_ascend_direct',
        'prefill', 'decode', 'dp_size', 'tp_size',
        'num_speculative_tokens', 'cudagraph_mode',
        'FULL_DECODE_ONLY', 'cudagraph_capture_sizes',
        'enable_npugraph_ex', 'fuse_muls_add',
        'multistream_overlap_shared_expert',
        'recompute_scheduler_enable', 'fuse_qknorm_rope',
        'MooncakeConnectorV1', 'kv_producer', 'kv_consumer',
        'do_remote_decode', 'do_remote_prefill',
        'remote_engine_id', 'remote_block_ids',
        'remote_host', 'remote_port', 'aborted_request',
        'stream_options', 'max_completion_tokens',
        'min_tokens', 'max_tokens',
        'X-Request-Id', 'Authorization',
        'OPENAI_API_KEY', 'base_delay', 'max_retries',
        'attempt', 'last_exc',
    ]

    looks_like_code = any(text.startswith(ind) or text.startswith('    ' + ind) for ind in code_indicators if len(ind) > 3)

    if looks_like_code or is_code_style(style_name):
        if not in_code_block:
            if any(kw in text for kw in ['import ', 'from ', 'def ', 'class ', '#!/']):
                code_lang = 'python'
            elif any(kw in text for kw in ['export ', 'nic_name=', 'local_ip=', 'vllm serve', 'docker ', 'hccn_tool', 'for i in', 'if [']):
                code_lang = 'bash'
            else:
                code_lang = ''
            output_lines.append(f'```{code_lang}')
            in_code_block = True
        output_lines.append(text)
    else:
        if in_code_block:
            output_lines.append('```')
            output_lines.append('')
            in_code_block = False
        output_lines.append(text)

    i += 1

if in_code_block:
    output_lines.append('```')
    output_lines.append('')
    in_code_block = False

output_lines.append('')
output_lines.append('')

tables = doc.tables
for ti, table in enumerate(tables):
    rows_data = []
    for ri, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace('\n', ' ').replace('\r', ' ').replace('|', '\\|')
            cells.append(cell_text)
        rows_data.append(cells)

    if not rows_data:
        continue

    ncols = max(len(r) for r in rows_data)
    header = rows_data[0] if rows_data else []
    while len(header) < ncols:
        header.append('')

    output_lines.append('')
    output_lines.append('| ' + ' | '.join(header) + ' |')
    output_lines.append('| ' + ' | '.join(['---'] * ncols) + ' |')
    for row in rows_data[1:]:
        while len(row) < ncols:
            row.append('')
        output_lines.append('| ' + ' | '.join(row) + ' |')
    output_lines.append('')

out_path = r'C:\Repo\vllm-ascend\docs\source\tutorials\features\extracted_content.txt'
with open(out_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(output_lines))

print(f'Content extracted to {out_path}')
print(f'Total lines: {len(output_lines)}')
