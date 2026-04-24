import sys
import os
import re
import json
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r'C:\Repo\vllm-ascend\docs\source\tutorials\features\GLM-5.1昇腾支持一本通-vllm-ascend v0.18.0版本配套-v1.0.docx'
doc = Document(doc_path)

def get_heading_level(style_name):
    mapping = {
        'Heading 1': 1, 'Heading 2': 2, 'Heading 3': 3,
        'Heading 4': 4, 'Heading 5': 5, 'Heading 6': 6,
    }
    return mapping.get(style_name, 0)

def is_list_style(style_name):
    return style_name in ['List Paragraph', 'List Bullet', 'List Number', 'List Continue',
                          'List 2', 'List 3', 'List 4', 'List 5']

def is_code_style(style_name):
    return style_name in ['Code', 'Source Code', 'No Spacing', 'Normal (Web)']

def looks_like_code(text):
    if not text:
        return False
    code_starts = [
        'export ', 'import ', 'from ', 'def ', 'class ', 'for i in', 'for j in',
        'if __', 'if [', 'if !', 'if $', 'while ', 'try:', 'except ', 'with ',
        'vllm serve', 'vllm ', 'docker run', 'docker pull', 'docker load',
        'docker exec', 'docker images', 'pip install', 'pip list',
        'hccn_tool', 'npu-smi', 'cat /', '#!/', 'nic_name=', 'local_ip=',
        'node0_ip=', 'IMAGES_ID=', 'NAME=', 'echo "error', 'exit 1',
        'fi', 'done', 'done;',
        'ASCEND_RT_VISIBLE_DEVICES', 'LD_LIBRARY_PATH=',
        'PYTORCH_NPU_ALLOC_CONF', 'HCCL_', 'OMP_', 'VLLM_',
        'ASCEND_', 'GLOO_SOCKET', 'TP_SOCKET', 'TASK_QUEUE',
        'ACL_OP_INIT', 'VLLM_NIXL', 'ASCEND_RT',
        'self.', 'return ', 'parser.add_argument', 'args =',
        'command =', 'subprocess.', 'process.', 'async def',
        'await ', 'httpx.', 'heapq.', 'logger.',
        'asyncio.', 'threading.', 'multiprocessing.',
        'response =', 'headers =', 'req_data', 'client.',
        'app =', 'proxy_state', 'global_args',
        'sys.exit', 'os.path', 'os.environ',
        'kv_connector', 'kv_role', 'kv_port', 'engine_id',
        'MooncakeConnectorV1', 'kv_producer', 'kv_consumer',
        'self_host', 'other_host',
    ]
    for start in code_starts:
        if text.startswith(start) or text.startswith('    ' + start) or text.startswith('\t' + start):
            return True
    vllm_flags = ['--host', '--port', '--data-parallel', '--tensor-parallel',
                  '--quantization', '--seed', '--served-model', '--max-num',
                  '--trust-remote', '--gpu-memory', '--enable-', '--async-',
                  '--compilation-config', '--speculative-config', '--additional-config',
                  '--kv-transfer-config', '--reasoning-parser', '--tool-call-parser',
                  '--enforce-eager', '--headless']
    for flag in vllm_flags:
        if text.startswith(flag):
            return True
    if text.startswith('--') and ('=' in text or "'" in text or '"' in text):
        return True
    return False

def detect_code_lang(text):
    if any(kw in text for kw in ['import ', 'from ', 'def ', 'class ', '#!/', 'self.', 'async def']):
        return 'python'
    if any(kw in text for kw in ['export ', 'nic_name=', 'local_ip=', 'vllm serve', 'docker ',
                                  'hccn_tool', 'for i in', 'if [', 'if !', '#!/bin/bash']):
        return 'bash'
    return ''

def get_para_indent_level(para):
    pf = para.paragraph_format
    if pf and pf.left_indent is not None:
        return int(pf.left_indent / 720)
    return 0

def get_text_with_indent(para):
    xml_elem = para._element
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    runs = xml_elem.findall('.//w:r', nsmap)
    if not runs:
        return para.text

    result_parts = []
    for run in runs:
        rpr = run.find('w:rPr', nsmap)
        indent_elem = None
        if rpr is not None:
            indent_elem = rpr.find('w:indent', nsmap)

        text_elems = run.findall('w:t', nsmap)
        run_text = ''.join(t.text or '' for t in text_elems)

        if indent_elem is not None:
            h_val = indent_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi')
            val = indent_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if val:
                try:
                    indent_chars = int(val) // 240
                    result_parts.append(' ' * indent_chars + run_text)
                    continue
                except:
                    pass

        result_parts.append(run_text)

    return ''.join(result_parts)

body = doc.element.body
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

all_elements = []
for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        all_elements.append(('para', child))
    elif tag == 'tbl':
        all_elements.append(('table', child))

para_map = {}
for para in doc.paragraphs:
    para_map[id(para._element)] = para

table_map = {}
for table in doc.tables:
    table_map[id(table._element)] = table

output = []
in_code = False
code_lang = ''
prev_was_heading = False

for elem_type, elem in all_elements:
    if elem_type == 'table':
        if id(elem) not in table_map:
            continue
        table = table_map[id(elem)]
        if in_code:
            output.append('```')
            output.append('')
            in_code = False

        rows_data = []
        for ri, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ').replace('\r', ' ').replace('|', '\\|')
                cells.append(cell_text)
            rows_data.append(cells)

        if not rows_data:
            continue

        ncols = max(len(r) for r in rows_data)
        header = rows_data[0]
        while len(header) < ncols:
            header.append('')

        if all(c.strip() == '' for c in header):
            continue

        output.append('')
        output.append('| ' + ' | '.join(header) + ' |')
        output.append('| ' + ' | '.join(['---'] * ncols) + ' |')
        for row in rows_data[1:]:
            while len(row) < ncols:
                row.append('')
            if all(c.strip() == '' for c in row):
                continue
            output.append('| ' + ' | '.join(row) + ' |')
        output.append('')
        prev_was_heading = False
        continue

    if elem_type == 'para':
        if id(elem) not in para_map:
            continue
        para = para_map[id(elem)]
        style = para.style.name if para.style else 'None'
        text = para.text.strip()

        if style.startswith('toc') or style.startswith('Table'):
            continue

        if text in ('目  录', '目录'):
            continue

        heading_level = get_heading_level(style)

        if heading_level > 0:
            if in_code:
                output.append('```')
                output.append('')
                in_code = False
            clean = re.sub(r'^\d+\.\d+\.\d+\.\d+\s+', '', text)
            clean = re.sub(r'^\d+\.\d+\.\d+\s+', '', clean)
            clean = re.sub(r'^\d+\.\d+\s+', '', clean)
            clean = re.sub(r'^\d+\s+', '', clean)
            output.append(f'{"#" * heading_level} {clean}')
            output.append('')
            prev_was_heading = True
            continue

        if not text:
            if not in_code and not prev_was_heading:
                output.append('')
            continue

        prev_was_heading = False

        if is_list_style(style):
            if in_code:
                output.append('```')
                output.append('')
                in_code = False
            level = get_para_indent_level(para)
            indent = '  ' * level
            clean = re.sub(r'^[\d]+\.\s*', '', text)
            output.append(f'{indent}- {clean}')
            continue

        is_code = looks_like_code(text) or is_code_style(style)

        if is_code:
            if not in_code:
                code_lang = detect_code_lang(text)
                output.append(f'```{code_lang}')
                in_code = True
            output.append(text)
        else:
            if in_code:
                output.append('```')
                output.append('')
                in_code = False
            output.append(text)

if in_code:
    output.append('```')
    output.append('')

out_path = r'C:\Repo\vllm-ascend\docs\source\tutorials\features\glm-5.1-ascend-support.md'
with open(out_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(output))

print(f'Markdown written to {out_path}')
print(f'Total lines: {len(output)}')
