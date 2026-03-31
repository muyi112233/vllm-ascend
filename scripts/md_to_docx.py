#!/usr/bin/env python3
"""Convert Markdown to Word document."""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def parse_markdown_to_word(md_file: str, output_file: str):
    """Parse markdown file and create Word document."""
    doc = Document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []
    code_block_lang = ""
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        if line.strip().startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_block_content)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
                code_block_lang = line.strip()[3:].strip()
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        if line.strip().startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            if len(table_rows) > 1:
                create_table(doc, table_rows)
            table_rows = []
            in_table = False
        
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        elif line.strip().startswith('> '):
            text = line.strip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            if text.startswith('**') and ':**' in text:
                match = re.match(r'\*\*(.+?):\*\*\s*(.*)', text)
                if match:
                    p.clear()
                    run1 = p.add_run(match.group(1) + ': ')
                    run1.bold = True
                    run1.italic = True
                    run2 = p.add_run(match.group(2))
                    run2.italic = True
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            p = doc.add_paragraph(text, style='List Bullet')
        elif line.strip().startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
        elif line.strip() == '---':
            doc.add_paragraph('─' * 50)
        elif line.strip():
            text = line
            p = doc.add_paragraph()
            add_formatted_text(p, text)
        
        i += 1
    
    if in_table and len(table_rows) > 1:
        create_table(doc, table_rows)
    
    doc.save(output_file)
    print(f"Word document saved to: {output_file}")


def create_table(doc, table_rows):
    """Create a table from markdown table rows."""
    rows = []
    for row in table_rows:
        if row.strip().startswith('|--') or row.strip().startswith('| --'):
            continue
        cells = [cell.strip() for cell in row.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_text
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    doc.add_paragraph()


def add_formatted_text(paragraph, text):
    """Add text with markdown formatting to paragraph."""
    parts = []
    current = ""
    i = 0
    
    while i < len(text):
        if text[i:i+2] == '**':
            if current:
                parts.append(('normal', current))
                current = ""
            j = text.find('**', i + 2)
            if j != -1:
                parts.append(('bold', text[i+2:j]))
                i = j + 2
            else:
                current += text[i:i+2]
                i += 2
        elif text[i] == '`':
            if current:
                parts.append(('normal', current))
                current = ""
            j = text.find('`', i + 1)
            if j != -1:
                parts.append(('code', text[i+1:j]))
                i = j + 1
            else:
                current += text[i]
                i += 1
        elif text[i] == '[':
            j = text.find(']', i)
            if j != -1 and j + 1 < len(text) and text[j+1] == '(':
                k = text.find(')', j + 2)
                if k != -1:
                    if current:
                        parts.append(('normal', current))
                        current = ""
                    link_text = text[i+1:j]
                    link_url = text[j+2:k]
                    parts.append(('link', link_text, link_url))
                    i = k + 1
                else:
                    current += text[i]
                    i += 1
            else:
                current += text[i]
                i += 1
        else:
            current += text[i]
            i += 1
    
    if current:
        parts.append(('normal', current))
    
    for part in parts:
        if part[0] == 'normal':
            run = paragraph.add_run(part[1])
        elif part[0] == 'bold':
            run = paragraph.add_run(part[1])
            run.bold = True
        elif part[0] == 'code':
            run = paragraph.add_run(part[1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif part[0] == 'link':
            run = paragraph.add_run(part[1])
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True


if __name__ == '__main__':
    md_file = r'C:\Repo\vllm-ascend\docs\source\tutorials\features\pd_colocated_yuanrong_glm47_single_instance_cn.md'
    output_file = r'C:\Repo\vllm-ascend\docs\source\tutorials\features\pd_colocated_yuanrong_glm47_single_instance_cn.docx'
    parse_markdown_to_word(md_file, output_file)
